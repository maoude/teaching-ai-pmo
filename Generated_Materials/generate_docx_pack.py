from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "PMO_AI_Module_Materials"
OUTPUT = ROOT / "Generated_Materials" / "PMO_AI_Module_Pack.docx"


def add_markdown_file(document: Document, path: Path) -> None:
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line[:3].replace(".", "").isdigit() and ". " in line[:5]:
            document.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif line.startswith("|"):
            document.add_paragraph(line)
        else:
            document.add_paragraph(line)


def main() -> None:
    document = Document()
    document.add_heading("AI-Assisted Gamified PMO Module Pack", level=0)
    document.add_paragraph("From Plan to Control: Managing a Project Like a Real PMO")

    files = [
        "01_Module_Overview.md",
        "02_Gamification_Design.md",
        "03_Instructor_Runbook.md",
        "04_Student_Worksheet.md",
        "05_Quiz_And_Exit_Ticket.md",
        "06_Assessment_Rubric.md",
        "07_Learning_Designer_Structure.md",
        "08_Explanatory_Video_Script.md",
    ]

    for index, filename in enumerate(files):
        if index:
            document.add_page_break()
        add_markdown_file(document, SOURCE / filename)

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
